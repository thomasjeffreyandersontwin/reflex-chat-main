

import pandas as pd
import glob
import os
import numpy as np
from ast import literal_eval
import tiktoken
from docx import Document

enter_char = "\n"


GLOBAL_ROOT_FOLDER = './data'

class Heading:
    ENC = tiktoken.encoding_for_model("gpt-4")
    _GLOBAL_HEADING_COLLAPSE_LEVEL = 3

    @classmethod
    def GLOBAL_HEADING_COLLAPSE_LEVEL(cls, value=None):
        """Get the class-level collapse level."""
        if value is not None:
            cls._GLOBAL_HEADING_COLLAPSE_LEVEL = value
        return cls._GLOBAL_HEADING_COLLAPSE_LEVEL

    _GLOBAL_MAX_TOKEN_LENGTH = 100000
    
    @classmethod
    def GLOBAL_MAX_TOKEN_LENGTH(cls, value=None):
        """Get the class-level collapse level."""
        if value is not None:
            cls._GLOBAL_MAX_TOKEN_LENGTH = value
        return cls._GLOBAL_MAX_TOKEN_LENGTH

    
    def __init__(self, title, level=1, parent=None):
        self.level = level
        self.title = title
        self.text = ""  # Start with empty text for non-heading content
        self._sub_headings = []
        self.parent = parent
        self.list_buffer = []  # To store and format list items before appending to text
        self.collapse = False  # Determines if sub-headings are collapsed
        self.tuning_examples = []
        self.embedding= None
    
    @property
    def text(self):
        if Heading.GLOBAL_HEADING_COLLAPSE_LEVEL() != 0 and self.level >= Heading.GLOBAL_HEADING_COLLAPSE_LEVEL():
            combined_text = self.full_text
            if len(Heading.ENC.encode(combined_text)) > Heading.GLOBAL_MAX_TOKEN_LENGTH():
                return self._text  # Return only this heading's text if combined text exceeds limit
            return combined_text
        return self._text   
    @text.setter
    def text(self, value):
        self._text = value

    @property
    def token_length(self):
        # Determine if this heading should use collapsed text or not
        if Heading.GLOBAL_HEADING_COLLAPSE_LEVEL() != 0 and self.level >= Heading.GLOBAL_HEADING_COLLAPSE_LEVEL():
            # Calculate the token length of the combined text
            combined_text = self.full_text
            combined_text_length = len(Heading.ENC.encode(combined_text))
            if combined_text_length > Heading.GLOBAL_MAX_TOKEN_LENGTH():
                # If the combined text exceeds the max token length, use only this heading's text
                return len(Heading.ENC.encode(self.full_title + " " + self._text))
            return combined_text_length
        # If not collapsing, only use this heading's text and title
        return len(Heading.ENC.encode(self.full_title + " " + self._text))

    def add_text(self, text, list_level=0):
        trimmed_text = text.strip()
        if trimmed_text:
            if list_level == 1:
                trimmed_text =  "\n"+ trimmed_text + "->"
            if list_level == 2:
                trimmed_text =  ", " + trimmed_text  + " "
            
            if not list_level:
                trimmed_text = '\n'+ trimmed_text 
                
            self.text += trimmed_text

    def finish_list(self, upto_level):
        # Finish the list up to the specified level
        while len(self.list_buffer) > upto_level:
            list_text = self.list_buffer.pop()
            if self.text and not self.text.endswith('\n'):
                self.text += '\n'
            self.text += list_text + '\n'

    @property
    def full_text(self):
   
        # Initially build full_text without checking sub_headings collapse condition
        #full_title = self.full_title
        text_with_subheadings =  self._text + '\n'
        # Check if sub_headings should be collapsed, if not, append their text
        for sub_heading in self._sub_headings:
            text_with_subheadings +=  "**"+  sub_heading.full_title +"**"+ sub_heading.full_text
        return text_with_subheadings
 
    @property
    def name(self):
        return SegmentableDocument.GLOBAL_ROOT_FOLDER() + "/" + self.title 
    
    @property
    def sub_headings(self):
        # Collapse sub_headings based on the encoded text length
        if Heading.GLOBAL_HEADING_COLLAPSE_LEVEL() != 0 and self.level >= Heading.GLOBAL_HEADING_COLLAPSE_LEVEL():
            if len(Heading.ENC.encode(self._text + ''.join(sub.full_text for sub in self._sub_headings))) <= Heading.GLOBAL_MAX_TOKEN_LENGTH():
                return []
        return self._sub_headings  # Directly use _sub_headings

    @property
    def full_title(self):
        if self.parent:
            return self.parent.full_title+"."+ self.title
        return self.title

    def __str__(self, level=0):
        indent = "    " * (self.level - 1)
        full_title = self.full_title
        ret = f"{indent}Title: {full_title}, Level: {self.level}, Text: {self.text}\n"
        if len(self.tuning_examples) > 0:
            backslash_char = "\\"
            enter_char = "\\n"
            # ret += f"{indent}Tuning Examples: {'\n'.join(self.tuning_examples)}\n\n"
            ret += f"{indent}Tuning Examples: {enter_char.join(self.tuning_examples)}{enter_char}{enter_char}"
        for sub_heading in self.sub_headings:
            ret += sub_heading.__str__(self.level)
        ret += "\n"
        return ret  
    
    def write_to_file(self, file):
        indent = "    " * (self.level - 1)
        file.write(f"{indent}Title: {self.full_title}  Level:{self.level} Token Count: {self.token_length}\n")
        file.write(f"{indent}Text: {self.text}\n" )
        
        # file.write(f"{indent}Tuning Examples: {'\n'.join(self.tuning_examples)} \n\n")
        file.write(f"{indent}Tuning Examples: {enter_char.join(self.tuning_examples)} {enter_char}{enter_char}")
        for sub_heading in self.sub_headings:
            sub_heading.write_to_file(file)
                     
    def collect_data(self, data=[]):
        data.append({
            'Full Title': self.full_title,
            'Text': self.text,
            'Token Length': self.token_length,
            'Embedding': self.embedding if self.embedding is not None else None,
            'Tuning Examples': '\n'.join(self.tuning_examples)
        })
        
        for sub_heading in self.sub_headings:
            sub_heading.collect_data(data)
        return data
    
    @property
    def headings_iterator(self):
        """Generator to iterate over all documents' headings."""
        for document in self.data_documents:
            yield from document.headings_iterator

    @property
    def nested_tuning_examples(self):
        """Collects all tuning examples in headings into a list."""
        examples = self.tuning_examples.copy()
        # Recursively collect examples from sub-headings
        for sub_heading in self._sub_headings:
            examples.extend(sub_heading.nested_tuning_examples)
        return examples
class SegmentableDocument:
    root_folder = './data'
        
    @classmethod
    def GLOBAL_ROOT_FOLDER(cls, value=None):
        if value is not None:
            cls.root_folder = value
        return cls.root_folder
    
    def __init__(self,name=""):
        self.headings = []
        if name:
            self.filepath = name
        else:
            self.filepath = None
        self.Document = None
    
    @property
    def name(self):
        return SegmentableDocument.GLOBAL_ROOT_FOLDER() +'/'+ self.filepath
    def load(self, filepath=""):
        if not filepath and self.filepath:
            filepath = self.filepath
        elif filepath:
            self.filepath = filepath
        elif not filepath and not self.filepath:
            raise Exception("No file path provided.")
        #test if  file is a path
        if  os.path.isfile(filepath +".docx"):
            self.Document = Document(filepath +'.docx')
        else:
            self.Document = Document(SegmentableDocument.GLOBAL_ROOT_FOLDER() + filepath +'.docx')
        current_heading = None
        last_was_bold = False  # Tracks the bold status of the last list item processed

        for para in self.Document.paragraphs:
            trimmed_text = para.text.strip()
            if not trimmed_text:
                continue  # Skip empty paragraphs

            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1])
                if not current_heading or level == 1:
                    current_heading = Heading(trimmed_text, level)
                    self.headings.append(current_heading)
                else:
                    while current_heading and current_heading.level >= level:
                        current_heading = current_heading.parent
                    new_heading = Heading(trimmed_text, level, parent=current_heading)
                    current_heading.sub_headings.append(new_heading)
                    current_heading = new_heading
            else:
                # Check if the paragraph is part of a list
                is_list = "list" in para.style.name.lower()
                if is_list:
                    bold = self.is_bold(para)
                    if bold:
                        list_level = 1  # Start of a new list or new top-level item
                    else:
                        list_level = 2  # Continuation of the previous list at level 2
                    last_was_bold = bold  # Update the bold tracking status
                    current_heading.add_text(trimmed_text, list_level)
                else:
                    
                    current_heading.add_text(trimmed_text)  # Add normal text
   
    def __str__(self):
        ret = "Document Headings:\n"
        for heading in self.headings:
            ret += heading.__str__(0)
        return ret
    
    @property
    def headings_iterator(self):
        """Generator to iterate over all documents' headings."""
        for document in self.data_documents:
            yield from document.headings_iterator
    
    @property
    def token_length(self):
        tokens = Heading.ENC.encode(self.full_text)
        return len(tokens)
    
    @property
    def tuning_examples(self):
        tuning_examples = []
        for heading in self.headings:
            tuning_examples.extend(heading.nested_tuning_examples)
        return tuning_examples
    @property
    def full_text(self):
        full_title = self.filepath
        ret = '***'+ full_title +"***"
        for heading in self.headings:
            ret += heading.full_text
        return ret
 
    def is_bold(self, para):
        """Determines if the majority of the text in a paragraph is bold."""
        bold_chars = sum(len(run.text) for run in para.runs if run.bold)
        total_chars = sum(len(run.text) for run in para.runs)
        return bold_chars > total_chars / 2  # More than half of the text is bold
    
    def get_list_level(self,para):
        if para.paragraph_format.left_indent:
            # Assuming a basic indentation of 0.5 inches (720 Twips) per list level
            indentation_twips = para.paragraph_format.left_indent.twips
            level = int(indentation_twips / 720)
            return level
        return 0  # No indentation implies the base level

    def save_to_file(self):
        filename = SegmentableDocument.GLOBAL_ROOT_FOLDER() + self.filepath +'_amazing_output.txt'
        with open(filename, 'w', encoding='utf-8') as file:
            for heading in self.headings:
                print(heading.embedding, ' embedding ', 'from ', self.filepath)
                heading.write_to_file(file)
    
    def remove_document(self, subFolder="."):
        filename = SegmentableDocument.GLOBAL_ROOT_FOLDER() + f'/{subFolder}/'+ self.filepath +'.csv'
        os.remove(filename)

    def save_to_csv(self, subFolder="."):
        
        # filename = self.filepath +'_amazing_output.csv'
        
        filename = SegmentableDocument.GLOBAL_ROOT_FOLDER() + f'/{subFolder}/'+ self.filepath +'.csv'
        print("saved to csv: ", filename)
        
        self.createDataFrame()
        self.dataFrame.to_csv(filename, index=False)

    def collect_data(self, data=[]):
        for heading in self.headings:
            data = heading.collect_data(data)
        return data
    
    def createDataFrame(self):
        data= []
        data = self.collect_data(data)
            
        self.dataFrame = pd.DataFrame(data, columns=['Full Title', 'Text', 'Token Length','Embedding'])
        print("created dataframe for ", self.filepath)
      
    @property
    def embeddings(self):
        embeddings = []
        for heading in self.headings_iterator:
            embeddings.append(heading.embedding)
        return self.document_source.embeddings
    
    def load_from_csv(self, aFileName="", subFolder="."):
        # if not aFileName:
        #     filename = SegmentableDocument.GLOBAL_ROOT_FOLDER() + '/output/' + self.filepath + '_output.csv'
        # else:
        #     filename = aFileName
        if not aFileName:
            filename = SegmentableDocument.GLOBAL_ROOT_FOLDER() + f'/{subFolder}/' + self.filepath + '.csv'
        else:
            filename = aFileName
        self.dataFrame = pd.read_csv(filename)
        current_heading = None
        headings_stack = []  # Use a stack to manage the current heading and its parents

        for _, row in self.dataFrame.iterrows():
            full_title = row['Full Title']
            if pd.notna(row['Embedding']):
                embedding = row['Embedding']  
                embedding = literal_eval(embedding)
                row['Embedding']  = embedding
            else:
                embedding = ""  

            title_parts = full_title.split('.')
            level = len(title_parts) - 1  # Level inferred by count of title parts
            title = title_parts[-1]  # Get the last part of the full title
            text = row['Text'] if pd.notna(row['Text']) else ""  # Handle NaN values by converting to empty string

            
            # Navigate up the stack to find the correct parent level
            while len(headings_stack) > level:
                headings_stack.pop()  # Pop until we find the parent level
            current_heading = headings_stack[-1] if headings_stack else None

            # Create a new heading
            new_heading = Heading(title, level + 1, parent=current_heading)
            new_heading.embedding = embedding 
            if current_heading:
                current_heading.sub_headings.append(new_heading)
            else:
                self.headings.append(new_heading)  # Add to the root level if no parent

            new_heading.text = text  # Set text directly
            headings_stack.append(new_heading)  # Push this heading onto the stack

    @property
    def headings_iterator(self):
        """Generates an iterator that walks through all headings, respecting max_token_length and collapse_level."""
        def walk(headings):
            for heading in headings:
                yield heading
                if heading.sub_headings and (Heading.GLOBAL_HEADING_COLLAPSE_LEVEL() == 0 or heading.level < Heading.GLOBAL_HEADING_COLLAPSE_LEVEL()):
                    if len(Heading.ENC.encode(heading.full_text)) <= Heading.GLOBAL_MAX_TOKEN_LENGTH():
                        yield from walk(heading.sub_headings)

        return walk(self.headings)
 
class DocumentFolder(SegmentableDocument):
    def __init__(self, path) -> None:
        super().__init__(path)
        #self.tuning_examples = []
        self.file_url=None
        self.data_documents = []
        self.folder_path = SegmentableDocument.GLOBAL_ROOT_FOLDER() +path
        print ( self.folder_path)
        if os.path.isdir(self.folder_path ):
            
            for file_path in glob.glob(os.path.join(self.folder_path, '*.csv')):
                file_name = os.path.basename(file_path).split('.')[0]
                print (file_name)
                document = SegmentableDocument(file_name)

                self.data_documents.append ( document)
        else:
            raise Exception("Invalid data path. Please provide a valid document object or directory containing DOCX files.")
    
    @property
    def name(self):
        return self.folder_path
             
    def load(self):
        for document in self.data_documents:
            document.load(self.folder_path + "/" + document.filepath)
    def save_to_file(self):
        for document in self.data_documents:
            document.save_to_file()

    def update_file_url(self, url: str):
        self.file_url = url
    
    def load_from_csv(self):
        if self.file_url is None:
            filename = self.folder_path +'/'+ 'content.csv'
            self.file_url = filename
        else:
            filename = self.file_url

        print("loading csv: ", filename)

        super().load_from_csv(filename)
        
    def collect_data(self, data=[]):
        
        for document in self.data_documents:
            data = document.collect_data(data)
        return data
    
    def save_to_csv(self):
        # filename = self.folder_path +'/'+ 'content.csv'
        filename = self.file_url
        
        data =[]
        for document in self.data_documents:
            data=document.collect_data(data)
        
        self.dataFrame = pd.DataFrame(data, columns=['Full Title', 'Text', 'Token Length','Embedding'])
        self.dataFrame.to_csv(filename, index=False)

    def get_document_by_name(self, document_name:str):
        """
        return SegementableDocument object by document name
        """
        for document in self.data_documents:
            if document.filepath == document_name:
                return document
        return None
    
    def get_document_list(self):
        """
        return an array of csv document file names
        """
        doc_list = []
        for document in self.data_documents:
            doc_list.append(document.filepath)
        return doc_list
    
    def create_document(self, document_name: str):
        document = SegmentableDocument(document_name)
        document.save_to_csv()
        self.data_documents.append(document)
        return document
    
    def create_chat_document(self, document_name: str):
        document = SegmentableDocument(document_name)
        document.save_to_csv(subFolder='chats')
        self.data_documents.append(document)
        return document


    @property
    def token_length(self):
        # Sum of token lengths of all documents
        return sum(doc.token_length for doc in self.data_documents)

    @property
    def full_text(self):
        # Concatenate full texts of all documents
        return '\n'.join(doc.full_text for doc in self.data_documents)


    def print_headings(self):
        for document in self.data_documents:
            document.print_headings(document.headings)   
            
    @property
    def headings_iterator(self):
        """Generator to iterate over all documents' headings."""
        for document in self.data_documents:
            yield from document.headings_iterator     
            
  

